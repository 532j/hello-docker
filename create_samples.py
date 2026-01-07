import os
from openpyxl import Workbook
from docx import Document

# 保存用フォルダの作成
os.makedirs("samples", exist_ok=True)

def create_files():
    for i in range(1, 11):
        # --- エクセルファイルの作成 ---
        wb = Workbook()
        ws = wb.active
        ws.title = f"Data_{i}"
        ws["A1"] = "ID"
        ws["B1"] = "内容"
        ws["A2"] = i
        ws["B2"] = f"これはサンプルエクセルファイル第 {i} 号のデータです。"
        wb.save(f"samples/sample_excel_{i}.xlsx")

        # --- ワードファイルの作成 ---
        doc = Document()
        doc.add_heading(f"サンプル文書 第 {i} 号", 0)
        doc.add_paragraph(f"これは練習用のワードファイルです。")
        doc.add_paragraph(f"整理番号: {i:03}")
        doc.add_paragraph(f"このテキストがAIによって正しく検索されるかテストしましょう。")
        doc.save(f"samples/sample_word_{i}.docx")

    print("✅ samples フォルダに20個のファイルを作成しました！")

if __name__ == "__main__":
    create_files()